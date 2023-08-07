import os
import googletrans
import textract
import PyPDF2
from reportlab.pdfgen import canvas
from docx import Document
from googletrans import Translator
traslator = Translator()

# 新建储存翻译文件的文件夹
def create_folder(folder_path):
    # 检查文件夹是否存在，如果不存在则创建它
    index = 1
    org_name = folder_path
    while True:
        try:
            os.mkdir(folder_path)
            print('')
            print("翻译后文件将储存在", folder_path)
            return folder_path
            break
        except FileExistsError:
            index += 1
            folder_path = org_name + '第' + str(index) + '版'

# 读取docx文件的功能         
def read_docx(file_path):
    doc = Document(file_path)
    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    return text
# 写入docx文件的功能
def write_to_docx(text, file_path):
    doc = Document()
    doc.add_paragraph(text)
    doc.save(file_path)

# 读取PDF文本的功能
def read_pdf(file_path):
    with open(file_path, 'rb') as file:
        reader = PyPDF2.PdfReader(file)
        text = ""
        for page in reader.pages:
            text += page.extract_text()
    return text
# 写入PDF文本的功能
def write_to_pdf(text, file_path):
    c = canvas.Canvas(file_path)
    c.drawString(100, 750, text)  # 设置文本起始位置和内容
    c.save()

# 处理路径相关
print('--------------------------------------------------------------------------------')
while True:
    try:
        path = input('翻译源文件夹（绝对）路径：')
        break
    except FileNotFoundError:
        print('找不到此目录，请重试')
print('')

# folder储存了该路径下所有的文件内容，为str类型，收集为list
folder = os.listdir(path)
for thing in folder:
    if not thing.endswith('.txt') and not thing.endswith('.docx') and not thing.endswith('.pdf'):
        folder.remove(thing)
print('路径中找到以下（可翻译的）文件：')
print(folder)

new_folder_path = create_folder(path + r'/翻译件')

 
# 选择源语言
print('')
print('请选择目标语言，单击 ENTER 将默认翻译至简体中文。\n1. 中文   2.法语   3.英语')
choice = input('请输入对应的数字：')

if choice == '2':
    lan_out = 'fr'
elif choice == '3':
    lan_out = 'en'
else:
    lan_out = 'zh-CN'
print('--------------------------------------------------------------------------------')
print('')

store = {}
for name in folder:
    try:
        doc = path + '/' + name
        if name.endswith('.txt'):
            file = open(doc,'r')
            file_content = file.read()
            
        elif name.endswith('.docx'):
            file_content = read_docx(doc)
        
        elif name.endswith('.pdf'):
            file_content = read_pdf(doc)
                 
        text = str(traslator.translate(file_content, dest = lan_out))
        store[name] = text
        file.close()
        
        print(name + '已被翻译')
        print('\n')
        
    except:
        print(name + ' 不是需要被翻译的文件，或不支持翻译')
        print('\n')
        
for item in store:
    write_path = new_folder_path + '/' + item
    if item.endswith('.txt'):
        with open(write_path,"w") as w:
            w.write(store[item])
    elif item.endswith('.docx'):
        write_to_docx(store[item],write_path)
    elif item.endswith('.pdf'):
        write_path += '.txt'
        with open(write_path,"w") as w:
            w.write(store[item])

print('--------------------------------------------------------------------------------')
print('翻译已完成')






















